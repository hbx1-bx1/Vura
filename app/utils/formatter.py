"""
VURA Report Formatter — Professional Report Generation & Export
═══════════════════════════════════════════════════════════════════
Handles CVE enrichment, Markdown/PDF/JSON export, script extraction,
executive reports, dual-report generation, and compliance mapping.
"""

import os
import re
import sys
import json
import datetime
import requests
from concurrent.futures import ThreadPoolExecutor
from rich.console import Console

console = Console()

# ✅ FIX #4 — مسار مطلق لمجلد التقارير
from pathlib import Path as _Path
_PROJECT_ROOT   = _Path(__file__).parent.parent.parent.absolute()
_REPORTS_ROOT   = str(_PROJECT_ROOT / "reports")


# ═══════════════════════════════════════════════════════════════════════════════
# COMPLIANCE MAPPING DATABASE
# ═══════════════════════════════════════════════════════════════════════════════
# ربط أنواع الثغرات بمعايير الامتثال — ISO 27001, NCA ECC, GDPR, PCI-DSS

COMPLIANCE_MAP = {
    # ── Access Control & Authentication ──────────────────────────────────────
    "default credentials": {
        "iso_27001": ["A.9.2.3 Management of privileged access rights", "A.9.4.3 Password management system"],
        "nca_ecc":   ["2-3-1 Identity and Access Management", "2-3-1-2 Privileged Access Management"],
        "pci_dss":   ["Req 2.1 Change vendor-supplied defaults", "Req 8.2 Unique identification"],
        "gdpr":      ["Art. 32 Security of processing"],
    },
    "weak password": {
        "iso_27001": ["A.9.4.3 Password management system", "A.9.2.1 User registration and de-registration"],
        "nca_ecc":   ["2-3-1 Identity and Access Management"],
        "pci_dss":   ["Req 8.2.3 Password complexity"],
        "gdpr":      ["Art. 32 Security of processing"],
    },
    "brute force": {
        "iso_27001": ["A.9.4.2 Secure log-on procedures", "A.9.4.3 Password management system"],
        "nca_ecc":   ["2-3-1 Identity and Access Management", "2-6-1 Threat Management"],
        "pci_dss":   ["Req 8.1.6 Lockout after failed attempts"],
        "gdpr":      ["Art. 32 Security of processing"],
    },
    "authentication bypass": {
        "iso_27001": ["A.9.4.2 Secure log-on procedures", "A.14.2.5 Secure system engineering principles"],
        "nca_ecc":   ["2-3-1 Identity and Access Management", "2-5-1 Application Security"],
        "pci_dss":   ["Req 6.5.10 Broken authentication"],
        "gdpr":      ["Art. 32 Security of processing", "Art. 25 Data protection by design"],
    },

    # ── Injection & Application Vulnerabilities ─────────────────────────────
    "sql injection": {
        "iso_27001": ["A.14.2.5 Secure system engineering principles", "A.14.1.2 Securing application services"],
        "nca_ecc":   ["2-5-1 Application Security", "2-5-1-1 Secure Software Development"],
        "pci_dss":   ["Req 6.5.1 Injection flaws"],
        "gdpr":      ["Art. 32 Security of processing", "Art. 25 Data protection by design"],
    },
    "xss": {
        "iso_27001": ["A.14.2.5 Secure system engineering principles"],
        "nca_ecc":   ["2-5-1 Application Security"],
        "pci_dss":   ["Req 6.5.7 Cross-site scripting"],
        "gdpr":      ["Art. 32 Security of processing"],
    },
    "cross-site scripting": {
        "iso_27001": ["A.14.2.5 Secure system engineering principles"],
        "nca_ecc":   ["2-5-1 Application Security"],
        "pci_dss":   ["Req 6.5.7 Cross-site scripting"],
        "gdpr":      ["Art. 32 Security of processing"],
    },
    "command injection": {
        "iso_27001": ["A.14.2.5 Secure system engineering principles", "A.12.4.1 Event logging"],
        "nca_ecc":   ["2-5-1 Application Security", "2-5-1-1 Secure Software Development"],
        "pci_dss":   ["Req 6.5.1 Injection flaws"],
        "gdpr":      ["Art. 32 Security of processing"],
    },
    "remote code execution": {
        "iso_27001": ["A.12.6.1 Management of technical vulnerabilities", "A.14.2.5 Secure system engineering"],
        "nca_ecc":   ["2-6-1 Threat Management", "2-4-1 Vulnerability Management"],
        "pci_dss":   ["Req 6.1 Identify security vulnerabilities", "Req 6.5.1 Injection flaws"],
        "gdpr":      ["Art. 32 Security of processing", "Art. 33 Notification of data breach"],
    },
    "rce": {
        "iso_27001": ["A.12.6.1 Management of technical vulnerabilities"],
        "nca_ecc":   ["2-6-1 Threat Management", "2-4-1 Vulnerability Management"],
        "pci_dss":   ["Req 6.1 Identify security vulnerabilities"],
        "gdpr":      ["Art. 32 Security of processing"],
    },

    # ── Network & Infrastructure ─────────────────────────────────────────────
    "open port": {
        "iso_27001": ["A.13.1.1 Network controls", "A.13.1.3 Segregation in networks"],
        "nca_ecc":   ["2-2-1 Network Security Management"],
        "pci_dss":   ["Req 1.1 Firewall configuration standards"],
        "gdpr":      ["Art. 32 Security of processing"],
    },
    "unencrypted": {
        "iso_27001": ["A.10.1.1 Policy on use of cryptographic controls", "A.14.1.2 Securing application services"],
        "nca_ecc":   ["2-7-1 Cryptography", "2-7-1-1 Encryption Standards"],
        "pci_dss":   ["Req 4.1 Strong cryptography for transmission"],
        "gdpr":      ["Art. 32(1)(a) Encryption of personal data"],
    },
    "ssl": {
        "iso_27001": ["A.10.1.1 Policy on use of cryptographic controls"],
        "nca_ecc":   ["2-7-1 Cryptography"],
        "pci_dss":   ["Req 4.1 Strong cryptography"],
        "gdpr":      ["Art. 32(1)(a) Encryption"],
    },
    "tls": {
        "iso_27001": ["A.10.1.1 Policy on use of cryptographic controls"],
        "nca_ecc":   ["2-7-1 Cryptography"],
        "pci_dss":   ["Req 4.1 Strong cryptography"],
        "gdpr":      ["Art. 32(1)(a) Encryption"],
    },
    "outdated software": {
        "iso_27001": ["A.12.6.1 Management of technical vulnerabilities", "A.12.5.1 Installation on operational systems"],
        "nca_ecc":   ["2-4-1 Vulnerability Management", "2-4-1-1 Patch Management"],
        "pci_dss":   ["Req 6.2 Security patches within one month"],
        "gdpr":      ["Art. 32 Security of processing"],
    },
    "missing patch": {
        "iso_27001": ["A.12.6.1 Management of technical vulnerabilities"],
        "nca_ecc":   ["2-4-1 Vulnerability Management", "2-4-1-1 Patch Management"],
        "pci_dss":   ["Req 6.2 Security patches"],
        "gdpr":      ["Art. 32 Security of processing"],
    },

    # ── Data Exposure ────────────────────────────────────────────────────────
    "data exposure": {
        "iso_27001": ["A.8.2.3 Handling of assets", "A.18.1.4 Privacy and protection of PII"],
        "nca_ecc":   ["2-9-1 Data Management and Privacy"],
        "pci_dss":   ["Req 3.4 Render PAN unreadable"],
        "gdpr":      ["Art. 5(1)(f) Integrity and confidentiality", "Art. 32 Security of processing"],
    },
    "information disclosure": {
        "iso_27001": ["A.8.2.3 Handling of assets", "A.14.2.5 Secure system engineering"],
        "nca_ecc":   ["2-9-1 Data Management and Privacy"],
        "pci_dss":   ["Req 6.5.6 Information leakage"],
        "gdpr":      ["Art. 5(1)(f) Integrity and confidentiality"],
    },
    "directory listing": {
        "iso_27001": ["A.14.2.5 Secure system engineering principles"],
        "nca_ecc":   ["2-5-1 Application Security"],
        "pci_dss":   ["Req 6.5.6 Information leakage"],
        "gdpr":      ["Art. 32 Security of processing"],
    },

    # ── Misconfiguration ─────────────────────────────────────────────────────
    "misconfiguration": {
        "iso_27001": ["A.14.2.5 Secure system engineering principles", "A.12.1.1 Documented operating procedures"],
        "nca_ecc":   ["2-1-1 Cybersecurity Governance", "2-5-1 Application Security"],
        "pci_dss":   ["Req 2.2 Configuration standards"],
        "gdpr":      ["Art. 32 Security of processing"],
    },
    "missing header": {
        "iso_27001": ["A.14.2.5 Secure system engineering principles"],
        "nca_ecc":   ["2-5-1 Application Security"],
        "pci_dss":   ["Req 6.5.10 Broken authentication and session management"],
        "gdpr":      ["Art. 25 Data protection by design"],
    },

    # ── Catch-all للـ CVEs العامة ──────────────────────────────────────────
    "cve-": {
        "iso_27001": ["A.12.6.1 Management of technical vulnerabilities"],
        "nca_ecc":   ["2-4-1 Vulnerability Management"],
        "pci_dss":   ["Req 6.1 Identify security vulnerabilities"],
        "gdpr":      ["Art. 32 Security of processing"],
    },
}


# ═══════════════════════════════════════════════════════════════════════════════
# EXISTING FUNCTIONS — لا تُلمس إلا بتحسين
# ═══════════════════════════════════════════════════════════════════════════════

def setup_directories():
    for sub in ["md", "pdf", "json", "sh", "docx"]:
        os.makedirs(os.path.join(_REPORTS_ROOT, sub), exist_ok=True)


def _fetch_cve(cve: str) -> str:
    """Fetch a single CVE from circl.lu and format it as a markdown list item."""
    try:
        resp = requests.get(f"https://cve.circl.lu/api/cve/{cve}", timeout=5)
        if resp.status_code == 200 and resp.json():
            data    = resp.json()
            cvss    = data.get("cvss", "N/A")
            summary = str(data.get("summary", "No description available."))[:250] + "..."
            return (
                f"- 🔴 **{cve}** | **CVSS Score:** `{cvss}` | "
                f"[View Official NVD Report](https://nvd.nist.gov/vuln/detail/{cve})\n"
                f"  - *{summary}*\n"
            )
    except Exception:
        pass
    return f"- 🔴 **{cve}** | [View Official NVD Report](https://nvd.nist.gov/vuln/detail/{cve})\n"


def enrich_with_cve_data(report_content):
    cves = sorted(set(re.findall(r"CVE-\d{4}-\d{4,7}", report_content.upper())))
    if not cves:
        return report_content

    console.print(f"\n[bold yellow][⚡] Radar Alert: Found {len(cves)} CVE(s). Fetching live threat intel...[/bold yellow]")

    # Fetch CVEs in parallel: 30 CVEs serial @ 5s timeout = 150s worst case;
    # parallel with 8 workers caps us at ~5s total even with every one timing out.
    with ThreadPoolExecutor(max_workers=min(8, len(cves))) as pool:
        lines = list(pool.map(_fetch_cve, cves))

    intel_section = "\n\n---\n\n## 🌍 VURA Live Threat Intelligence (NVD)\n" + "".join(lines)
    return report_content + intel_section


def generate_patch_script(report_content, session_id, approach="defense"):
    """
    ✅ FIX #3 — استبدال الـ Blacklist بـ Human-Review Gate حقيقي.

    المنطق الجديد:
    - يُحفظ السكربت بصلاحية 0o444 (للقراءة فقط، غير قابل للتنفيذ).
    - لا يمكن تشغيله إلا بعد أن يراجعه المستخدم يدوياً ويكتب: chmod +x <script>
    - هذا يُلغي الحاجة لـ Blacklist لأن المستخدم نفسه هو خط الدفاع الأخير.
    """
    # Accept ```bash / ```sh / ```shell fences — AI often picks any of these.
    bash_blocks = re.findall(
        r"```(?:bash|sh|shell)\s*\n?(.*?)```",
        report_content, re.DOTALL | re.IGNORECASE,
    )
    if not bash_blocks:
        return None

    all_scripts = "\n\n".join(block.strip() for block in bash_blocks)
    script_path = os.path.join(_REPORTS_ROOT, "sh", f"{session_id}.sh")

    title = "OFFENSIVE EXPLOIT SCRIPT" if approach == "offense" else "AUTO-GENERATED PATCH SCRIPT"
    safety_header = (
        f"#!/bin/bash\n"
        f"# ==========================================\n"
        f"# {title} BY VURA\n"
        f"# ==========================================\n"
        f"# ⚠️  WARNING: AI-GENERATED CODE.\n"
        f"# THIS FILE IS READ-ONLY BY DEFAULT.\n"
        f"# Review every command, then enable with:\n"
        f"#   chmod +x {os.path.basename(script_path)}\n"
        f"#   sudo ./{os.path.basename(script_path)}\n"
        f"# ==========================================\n"
        f"set -euo pipefail\n\n"
    )

    try:
        os.makedirs(os.path.dirname(script_path), exist_ok=True)
        with open(script_path, "w", encoding="utf-8") as f:
            if "#!/bin/bash" not in all_scripts:
                f.write(safety_header)
            else:
                all_scripts = all_scripts.replace("#!/bin/bash", safety_header.rstrip(), 1)
            f.write(all_scripts)

        # ✅ FIX #3 — 0o444 read-only on POSIX. Windows ignores POSIX bits;
        # use the read-only attribute there instead so the safety gate still applies.
        if sys.platform.startswith("win"):
            try:
                import stat
                os.chmod(script_path, stat.S_IREAD)
            except OSError:
                pass
        else:
            os.chmod(script_path, 0o444)
        return script_path
    except Exception as e:
        console.print(f"[dim red][!] Script generation failed: {e}[/dim red]")
        return None


def save_markdown_report(report_content, session_id, approach="defense"):
    setup_directories()
    enriched_content = enrich_with_cve_data(report_content)
    filename = os.path.join(_REPORTS_ROOT, "md", f"{session_id}.md")
    try:
        with open(filename, "w", encoding="utf-8") as f:
            f.write(enriched_content)
        patcher_script = generate_patch_script(enriched_content, session_id, approach)
        return filename, patcher_script, enriched_content
    except Exception as e:
        console.print(f"[bold red][!] Failed to save Markdown: {e}[/bold red]")
        return None, None, enriched_content


def save_json_report(json_content, session_id):
    setup_directories()
    filename = os.path.join(_REPORTS_ROOT, "json", f"{session_id}.json")
    try:
        clean_content = json_content.replace("```json", "").replace("```", "").strip()
        parsed = json.loads(clean_content)
        with open(filename, "w", encoding="utf-8") as f:
            json.dump(parsed, f, indent=4, ensure_ascii=False)
        return filename
    except Exception:
        with open(filename, "w", encoding="utf-8") as f:
            f.write(json_content)
        return filename


def export_to_pdf(markdown_content, session_id, analyst_name="", company_name="",
                  company_logo_path=""):
    """
    تصدير تقرير PDF مع branding اختياري.

    Parameters:
        analyst_name     : اسم المحلل — يُعرض في header
        company_name     : اسم الشركة — يُعرض في header
        company_logo_path: مسار لوجو الشركة (PNG/JPG)
    """
    setup_directories()

    try:
        import markdown
        from weasyprint import HTML
    except ImportError as e:
        console.print(f"[dim red][!] PDF export requires: pip install weasyprint markdown[/dim red]")
        console.print(f"[dim red]    On Mac: brew install pango gobject-introspection[/dim red]")
        return None

    # ── Branding header ──
    branding_html = ""
    if company_name or analyst_name:
        logo_tag = ""
        if company_logo_path and os.path.exists(company_logo_path):
            import base64
            with open(company_logo_path, "rb") as img:
                b64 = base64.b64encode(img.read()).decode()
            ext = company_logo_path.rsplit(".", 1)[-1].lower()
            mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg"}.get(ext, "image/png")
            logo_tag = f'<img src="data:{mime};base64,{b64}" style="height:50px;margin-right:15px;vertical-align:middle;">'

        branding_parts = []
        if company_name:
            branding_parts.append(f"<strong>{company_name}</strong>")
        if analyst_name:
            branding_parts.append(f"Analyst: {analyst_name}")

        branding_html = (
            f'<div style="text-align:center;padding:15px 0;border-bottom:2px solid #34495e;margin-bottom:20px;">'
            f'{logo_tag}{" | ".join(branding_parts)}'
            f'</div>'
        )

    # NOTE: do NOT @import Google Fonts here — weasyprint would fetch over the
    # network during PDF generation and hang offline / on slow DNS. Use only
    # locally available system fonts; if Tajawal is installed it will still
    # resolve, otherwise the fallback chain kicks in.
    css_style = """
    @page { size: A4; margin: 2cm; }
    body { font-family: 'Tajawal', 'Noto Sans Arabic', 'DejaVu Sans', Arial, sans-serif; line-height: 1.6; color: #2c3e50; }
    h1 { color: #1abc9c; border-bottom: 2px solid #34495e; padding-bottom: 10px; text-align: center; }
    h2 { color: #2980b9; margin-top: 30px; border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; }
    pre { background-color: #2c3e50; color: #ecf0f1; padding: 15px; border-radius: 8px; direction: ltr; text-align: left; overflow-x: auto; }
    code { font-family: Consolas, monospace; background-color: #ecf0f1; padding: 2px 5px; border-radius: 3px; color: #c0392b; direction: ltr; }
    pre code { background-color: transparent; color: #ecf0f1; padding: 0; }
    table { border-collapse: collapse; width: 100%; margin-top: 15px; margin-bottom: 15px; }
    th, td { border: 1px solid #bdc3c7; padding: 12px; }
    th { background-color: #34495e; color: white; }
    tr:nth-child(even) { background-color: #f9f9f9; }
    .footer { text-align: center; font-size: 0.8em; color: #95a5a6; margin-top: 30px; padding-top: 10px; border-top: 1px solid #bdc3c7; }
    """
    html_content = markdown.markdown(markdown_content, extensions=["tables", "fenced_code"])

    footer = '<div class="footer">Generated by VURA — Vulnerability Reporting AI</div>'

    full_html = (
        f"<!DOCTYPE html><html dir='auto'>"
        f"<head><meta charset='utf-8'><style>{css_style}</style></head>"
        f"<body>{branding_html}{html_content}{footer}</body></html>"
    )
    pdf_filename = os.path.join(_REPORTS_ROOT, "pdf", f"{session_id}.pdf")
    try:
        HTML(string=full_html).write_pdf(pdf_filename)
        return pdf_filename
    except (OSError, IOError) as e:
        # Disk full / permission issues — the user needs to see these.
        console.print(f"[bold red][!] PDF write failed (filesystem): {e}[/bold red]")
        return None
    except Exception as e:
        # weasyprint raises subclasses of Exception for network-font timeouts,
        # malformed HTML, missing Pango libs, etc. Don't swallow silently —
        # surface the error class + message so the user can act on it.
        console.print(
            f"[bold red][!] PDF generation failed: {type(e).__name__}: {e}[/bold red]"
        )
        console.print(
            "[dim]    If this mentions fonts or network, the report has been saved "
            "as Markdown — re-export PDF when the issue is resolved.[/dim]"
        )
        return None


def export_to_docx(markdown_content, session_id, analyst_name="", company_name="",
                   company_logo_path=""):
    """
    تصدير تقرير DOCX قابل للتعديل — خاص بخطة Max.

    Parameters:
        markdown_content  : محتوى التقرير (Markdown)
        session_id        : معرّف الجلسة
        analyst_name      : اسم المحلل
        company_name      : اسم الشركة
        company_logo_path : مسار لوجو الشركة (PNG/JPG)

    Returns:
        str : مسار ملف DOCX أو None
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        console.print("[dim red][!] python-docx not installed. Run: pip install python-docx[/dim red]")
        return None

    setup_directories()
    doc = DocxDocument()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Branding header ──
    if company_logo_path and os.path.exists(company_logo_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(company_logo_path, width=Inches(2))
        except Exception:
            pass

    if company_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(company_name)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    if analyst_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Analyst: {analyst_name}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    if company_name or analyst_name:
        doc.add_paragraph("─" * 60)

    # ── Parse Markdown → DOCX ──
    lines = markdown_content.split("\n")
    in_code_block = False
    code_buffer = []

    for line in lines:
        # ── Code blocks ──
        if line.strip().startswith("```"):
            if in_code_block:
                # نهاية code block
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                p.paragraph_format.left_indent = Cm(1)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        stripped = line.strip()

        # ── Headings ──
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)

        # ── Horizontal rule ──
        elif stripped.startswith("---") or stripped.startswith("═"):
            doc.add_paragraph("─" * 60)

        # ── Bullet points ──
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # remove bold markers
            text = re.sub(r'`(.+?)`', r'\1', text)          # remove code markers
            doc.add_paragraph(text, style="List Bullet")

        # ── Table rows (basic) ──
        elif stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells and not all(c.replace("-", "").replace(":", "") == "" for c in cells):
                p = doc.add_paragraph()
                text = "  |  ".join(cells)
                run = p.add_run(text)
                run.font.size = Pt(9)

        # ── Blockquote ──
        elif stripped.startswith("> "):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

        # ── Empty line ──
        elif not stripped:
            continue

        # ── Normal paragraph ──
        else:
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            text = re.sub(r'`(.+?)`', r'\1', text)
            doc.add_paragraph(text)

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated by VURA — Vulnerability Reporting AI")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    # ── Save ──
    docx_filename = os.path.join(_REPORTS_ROOT, "docx", f"{session_id}.docx")
    try:
        doc.save(docx_filename)
        return docx_filename
    except Exception as e:
        console.print(f"[dim red][!] DOCX export failed: {e}[/dim red]")
        return None


# ═══════════════════════════════════════════════════════════════════════════════
# NEW — COMPLIANCE MAPPING
# ═══════════════════════════════════════════════════════════════════════════════

def _detect_compliance_keywords(report_content):
    """
    فحص التقرير واستخراج الـ keywords التي تطابق COMPLIANCE_MAP.
    يرجع list من المفاتيح المطابقة.
    """
    content_lower = report_content.lower()
    matched_keys = []

    for keyword in COMPLIANCE_MAP:
        if keyword in content_lower:
            matched_keys.append(keyword)

    return matched_keys


def add_compliance_section(report_content, frameworks=None):
    """
    إضافة قسم Compliance Mapping لتقرير موجود.

    يفحص محتوى التقرير ويكتشف أنواع الثغرات تلقائياً،
    ثم يربطها بمعايير ISO 27001, NCA ECC, PCI-DSS, GDPR.

    Parameters:
        report_content : محتوى التقرير (Markdown string)
        frameworks     : قائمة المعايير المطلوبة (None = كلها)
                         مثال: ["iso_27001", "nca_ecc"]

    Returns:
        str : التقرير الأصلي + قسم Compliance مُضاف في النهاية
    """
    if not report_content:
        return report_content

    matched_keys = _detect_compliance_keywords(report_content)

    if not matched_keys:
        console.print("[dim yellow][~] VURA Compliance: No matching vulnerability patterns found.[/dim yellow]")
        return report_content

    console.print(f"[bold cyan][📋] VURA Compliance: Mapping {len(matched_keys)} finding(s) to regulatory frameworks...[/bold cyan]")

    # ── المعايير المطلوبة ──
    all_frameworks = ["iso_27001", "nca_ecc", "pci_dss", "gdpr"]
    active_frameworks = frameworks if frameworks else all_frameworks

    framework_labels = {
        "iso_27001": "ISO 27001:2022",
        "nca_ecc":   "NCA ECC (Saudi Arabia)",
        "pci_dss":   "PCI-DSS v4.0",
        "gdpr":      "GDPR (EU)",
    }

    # ── تجميع كل الـ controls المطابقة حسب المعيار ──
    framework_controls = {fw: set() for fw in active_frameworks}

    for keyword in matched_keys:
        mapping = COMPLIANCE_MAP.get(keyword, {})
        for fw in active_frameworks:
            controls = mapping.get(fw, [])
            for ctrl in controls:
                framework_controls[fw].add(ctrl)

    # ── بناء القسم ──
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    section = "\n\n---\n\n"
    section += "## 📋 VURA Compliance Mapping\n\n"
    section += f"*Auto-generated by VURA at {timestamp}*\n\n"
    section += f"**Findings matched:** {len(matched_keys)} vulnerability pattern(s)\n\n"

    for fw in active_frameworks:
        controls = sorted(framework_controls[fw])
        if not controls:
            continue

        label = framework_labels.get(fw, fw)
        section += f"### {label}\n\n"
        section += "| # | Control / Article | Status |\n"
        section += "|---|---|---|\n"
        for idx, ctrl in enumerate(controls, 1):
            section += f"| {idx} | {ctrl} | ⚠️ Potential Non-Compliance |\n"
        section += "\n"

    # ── ملاحظة قانونية ──
    section += (
        "> **Disclaimer:** This compliance mapping is AI-assisted and based on pattern matching. "
        "It does NOT constitute a formal audit or certification assessment. "
        "A qualified auditor should validate these findings against your organization's specific context.\n"
    )

    return report_content + section


# ═══════════════════════════════════════════════════════════════════════════════
# NEW — EXECUTIVE REPORT GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

def generate_executive_report(raw_data, session_id, language="English"):
    """
    توليد تقرير تنفيذي (Executive Summary) باستخدام AI.
    يستدعي ai_engine مع scan_type="executive".

    Parameters:
        raw_data   : البيانات الخام أو محتوى التقرير التقني
        session_id : معرّف الجلسة
        language   : لغة التقرير

    Returns:
        tuple: (md_path, pdf_path, content)
            - md_path  : مسار ملف MD المحفوظ (أو None)
            - pdf_path : مسار ملف PDF المحفوظ (أو None)
            - content  : محتوى التقرير التنفيذي
    """
    from app.core.ai_engine import generate_report

    console.print(f"\n[bold magenta][~] VURA: Generating Executive Summary ({language})...[/bold magenta]")

    exec_content = generate_report(
        raw_data,
        language=language,
        output_format="md",
        approach="defense",
        include_script=False,
        scan_type="executive",
    )

    if not exec_content or exec_content.startswith("# Connection Error") or exec_content.startswith("# Error"):
        console.print(f"[bold red][!] Executive report generation failed.[/bold red]")
        return None, None, exec_content

    # ── إضافة Compliance ──
    exec_content = add_compliance_section(exec_content)

    # ── حفظ MD ──
    setup_directories()
    exec_session = f"{session_id}_EXECUTIVE"
    md_path = os.path.join(_REPORTS_ROOT, "md", f"{exec_session}.md")

    try:
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(exec_content)
        console.print(f"[bold green][✔] Executive MD: ./{md_path}[/bold green]")
    except Exception as e:
        console.print(f"[bold red][!] Failed to save Executive MD: {e}[/bold red]")
        md_path = None

    # ── حفظ PDF ──
    pdf_path = export_to_pdf(exec_content, exec_session)
    if pdf_path:
        console.print(f"[bold green][✔] Executive PDF: ./{pdf_path}[/bold green]")

    return md_path, pdf_path, exec_content


# ═══════════════════════════════════════════════════════════════════════════════
# NEW — DUAL REPORT GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

def generate_dual_reports(raw_data, session_id, approach="defense",
                          language="English", output_format="md",
                          include_script=True, notify=None):
    """
    توليد تقريرين معاً: تقني كامل + تنفيذي للإدارة.

    التقرير التقني: يمر بالمسار العادي (CVE enrichment + script + compliance)
    التقرير التنفيذي: يُولّد بـ scan_type="executive" بدون scripts أو CVEs

    Parameters:
        raw_data       : البيانات الخام
        session_id     : معرّف الجلسة
        approach       : offense أو defense
        language       : لغة التقارير
        output_format  : md أو pdf (يُطبّق على التقني — التنفيذي دائماً PDF+MD)
        include_script : سكربت في التقرير التقني
        notify         : وضع Telegram (short/long أو None)

    Returns:
        dict: {
            "technical": {"md": path, "pdf": path, "content": str, "script": path},
            "executive": {"md": path, "pdf": path, "content": str},
            "compliance_added": bool,
        }
    """
    from app.core.ai_engine import generate_report

    console.print(f"\n[bold green]{'═' * 60}[/bold green]")
    console.print(f"[bold green]  VURA Dual Report Generation[/bold green]")
    console.print(f"[bold green]{'═' * 60}[/bold green]\n")

    results = {
        "technical":  {"md": None, "pdf": None, "content": None, "script": None},
        "executive":  {"md": None, "pdf": None, "content": None},
        "compliance_added": False,
    }

    # ══════════════════════════════════════════════════════════════════════════
    # 1. TECHNICAL REPORT
    # ══════════════════════════════════════════════════════════════════════════
    console.print("[bold cyan][1/2] Generating Technical Report...[/bold cyan]")

    tech_content = generate_report(
        raw_data,
        language=language,
        output_format="md",
        approach=approach,
        include_script=include_script,
        scan_type="terminal",
    )

    if tech_content and not tech_content.startswith("# Connection Error") and not tech_content.startswith("# Error"):
        # ── CVE Enrichment ──
        tech_enriched = enrich_with_cve_data(tech_content)

        # ── Compliance Mapping ──
        tech_enriched = add_compliance_section(tech_enriched)
        results["compliance_added"] = True

        # ── حفظ MD ──
        tech_md_path = os.path.join(_REPORTS_ROOT, "md", f"{session_id}.md")
        try:
            setup_directories()
            with open(tech_md_path, "w", encoding="utf-8") as f:
                f.write(tech_enriched)
            results["technical"]["md"] = tech_md_path
            console.print(f"[green]  ✓ Technical MD: ./{tech_md_path}[/green]")
        except Exception as e:
            console.print(f"[red]  ✘ Failed to save Technical MD: {e}[/red]")

        # ── Script ──
        if include_script:
            script_path = generate_patch_script(tech_enriched, session_id, approach)
            results["technical"]["script"] = script_path
            if script_path:
                stype = "EXPLOIT" if approach == "offense" else "PATCH"
                console.print(f"[green]  ✓ {stype} Script: ./{script_path}[/green]")

        # ── PDF إذا طُلب ──
        if output_format == "pdf":
            pdf_path = export_to_pdf(tech_enriched, session_id)
            results["technical"]["pdf"] = pdf_path
            if pdf_path:
                console.print(f"[green]  ✓ Technical PDF: ./{pdf_path}[/green]")

        results["technical"]["content"] = tech_enriched
    else:
        console.print(f"[bold red]  ✘ Technical report generation failed.[/bold red]")
        results["technical"]["content"] = tech_content

    # ══════════════════════════════════════════════════════════════════════════
    # 2. EXECUTIVE REPORT
    # ══════════════════════════════════════════════════════════════════════════
    console.print(f"\n[bold cyan][2/2] Generating Executive Summary...[/bold cyan]")

    exec_md, exec_pdf, exec_content = generate_executive_report(
        raw_data, session_id, language=language
    )

    results["executive"]["md"]      = exec_md
    results["executive"]["pdf"]     = exec_pdf
    results["executive"]["content"] = exec_content

    # ══════════════════════════════════════════════════════════════════════════
    # SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    console.print(f"\n[bold green]{'─' * 60}[/bold green]")
    console.print(f"[bold green]  ✓ Dual Report Generation Complete[/bold green]")

    files_generated = sum(1 for v in [
        results["technical"]["md"], results["technical"]["pdf"],
        results["technical"]["script"],
        results["executive"]["md"], results["executive"]["pdf"],
    ] if v is not None)

    console.print(f"[bold green]  Files generated: {files_generated}[/bold green]")
    console.print(f"[bold green]{'─' * 60}[/bold green]\n")

    # ── Telegram ──
    if notify and results["technical"]["content"]:
        try:
            from app.utils.notifier import send_telegram_alert
            path = results["technical"]["md"] or results["technical"]["pdf"] or "dual_report"
            send_telegram_alert(path, results["technical"]["content"], mode=notify)
        except Exception as e:
            console.print(f"[dim red][!] Telegram notification failed: {e}[/dim red]")

    return results
